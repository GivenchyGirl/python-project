from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)

document = Document()

# profile picture
document.add_picture('2.png', width=Inches(2.0))

# name phone and email details
print('loading...')
speak('Hello. I am Alice, your CV assistant.')
speak('And what is your name?')
name = input('What is your name? ')
speak('Hello ' + name + '. How are you today?')
speak('Please type your phone number.')
phone = input('What is your phone number? ')
speak('Please type your email address.')
email = input('What is your email? ')

document.add_paragraph(
    name + ' / ' + phone + ' / ' + email)

# about me
document.add_heading('About me')
speak('Please tell us about yourself')
document.add_paragraph(
    input('Tell us about yourself? ')
)

# work experience
document.add_heading('Work Experience')
p = document.add_paragraph ()
speak('Let us talk about your work experience. Please enter your company names and dates of work.')
company = input('Enter company ')
from_date = input('From Date ')
to_date = input('To Date ')

p.add_run(company + ' ').bold = True
p.add_run(from_date + '- ' + to_date + '\n').italic = True

speak('Describe your experience at ' + company)
experience_details = input(
    'Describe your experience at ' + company + ' '
)
p.add_run(experience_details)


# more experiences
speak('Have you had more work experiences?')
while True:
    has_more_experiences = input('Do you have more experiences? Yes or No ')
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()
        company = input('Enter company ')
        from_date = input('From Date ')
        to_date = input('To Date ')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + '- ' + to_date + '\n').italic = True

        experience_details = input(
            'Describe your experience at ' + company + ' ')
        p.add_run(experience_details)
    else:
        break

# skills
speak('Let us talk about your skills. Please list the most important from your point of view.')
document.add_heading('Skills')
pp = document.add_paragraph ()
skill = input('What are your skills? ')
pp.add_run(skill)
pp.style = 'List Bullet'

while True:
    speak('Anything else?')
    has_more_skills = input('List one more skill or type "No more" ')
    if has_more_skills == 'No more':
        break
    else:
        pp = document.add_paragraph(has_more_skills)
        pp.style = 'List Bullet'

# footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV generated using Python Tutorial for Beginners course by Amigoscode"

document.save('cv.docx')
print('saved')
speak('Thank you. Your CV is saved. Have a nice day.')
